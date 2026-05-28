from __future__ import annotations

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# Base Loader 
class BaseLoader(ABC):
    @abstractmethod
    def load(self, source: str) -> list[dict]:
        """Return list of page dicts: {text, page_num, source_path}."""


# PDF Loader 
class PDFLoader(BaseLoader):
    """Loads PDF files using pdfplumber (handles tables, multi-column layouts)."""

    def load(self, source: str) -> list[dict]:
        import pdfplumber

        pages = []
        with pdfplumber.open(source) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(
                        {"text": text, "page_num": i, "source_path": source}
                    )
        logger.info(f"PDF loaded: {source} ({len(pages)} pages)")
        return pages


# DOCX Loader
class DOCXLoader(BaseLoader):
    """Loads DOCX files preserving heading structure."""

    def load(self, source: str) -> list[dict]:
        from docx import Document

        doc = Document(source)
        sections: list[dict] = []
        current_heading = ""
        current_text: list[str] = []

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading"):
                if current_text:
                    sections.append(
                        {
                            "text": f"{current_heading}\n\n" + "\n".join(current_text),
                            "page_num": 1,
                            "source_path": source,
                            "section": current_heading,
                        }
                    )
                current_heading = text
                current_text = []
            else:
                current_text.append(text)

        if current_text:
            sections.append(
                {
                    "text": f"{current_heading}\n\n" + "\n".join(current_text),
                    "page_num": 1,
                    "source_path": source,
                    "section": current_heading,
                }
            )

        # Fall back to full text if no structure detected
        if not sections:
            full_text = "\n".join(
                p.text for p in doc.paragraphs if p.text.strip()
            )
            sections = [{"text": full_text, "page_num": 1, "source_path": source}]

        logger.info(f"DOCX loaded: {source} ({len(sections)} sections)")
        return sections


# Confluence Loader

class ConfluenceLoader(BaseLoader):
    """Loads pages from a Confluence space via the REST API."""

    def __init__(
        self,
        url: str,
        username: str,
        api_token: str,
        space_key: str,
    ):
        self.base_url = url.rstrip("/")
        self.auth = (username, api_token)
        self.space_key = space_key

    def load(self, source: str = "") -> list[dict]:
        """source is unused; loads all pages from the configured space."""
        import requests
        from markdownify import markdownify

        pages = []
        start = 0
        limit = 50

        while True:
            resp = requests.get(
                f"{self.base_url}/rest/api/content",
                auth=self.auth,
                params={
                    "spaceKey": self.space_key,
                    "type": "page",
                    "expand": "body.storage,metadata.labels,version",
                    "start": start,
                    "limit": limit,
                },
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            results = data.get("results", [])

            for page in results:
                html = page.get("body", {}).get("storage", {}).get("value", "")
                text = markdownify(html)
                pages.append(
                    {
                        "text": text,
                        "page_num": 1,
                        "source_path": f"{self.base_url}/wiki/spaces/{self.space_key}/pages/{page['id']}",
                        "title": page.get("title", ""),
                        "author": page.get("version", {}).get("by", {}).get("displayName"),
                        "last_updated": page.get("version", {}).get("when"),
                    }
                )

            if len(results) < limit:
                break
            start += limit

        logger.info(f"Confluence loaded: {len(pages)} pages from space {self.space_key}")
        return pages


# Wiki / HTML Loader

class WikiHTMLLoader(BaseLoader):
    """Loads internal wiki pages from an HTML file or URL."""

    def load(self, source: str) -> list[dict]:
        from bs4 import BeautifulSoup
        from markdownify import markdownify

        if source.startswith("http://") or source.startswith("https://"):
            import requests
            resp = requests.get(source, timeout=30)
            resp.raise_for_status()
            html = resp.text
        else:
            html = Path(source).read_text(encoding="utf-8")

        soup = BeautifulSoup(html, "lxml")
        # Remove nav, footer, scripts
        for tag in soup(["nav", "footer", "script", "style", "header"]):
            tag.decompose()

        content = soup.find("main") or soup.find("article") or soup.find("body")
        text = markdownify(str(content)) if content else ""

        return [{"text": text, "page_num": 1, "source_path": source}]


# Loader Factory

def get_loader(document_type: str, **kwargs) -> BaseLoader:
    loaders = {
        "pdf": PDFLoader,
        "docx": DOCXLoader,
        "wiki": WikiHTMLLoader,
        "html": WikiHTMLLoader,
    }
    if document_type == "confluence":
        return ConfluenceLoader(**kwargs)
    cls = loaders.get(document_type.lower())
    if cls is None:
        raise ValueError(f"Unsupported document type: {document_type}")
    return cls()
